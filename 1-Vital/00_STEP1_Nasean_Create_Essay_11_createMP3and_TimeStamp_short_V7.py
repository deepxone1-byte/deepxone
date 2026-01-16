# --- Start processing loop ---
import os
import re
import time
import csv
import docx
import requests
import whisper
from pathlib import Path
from datetime import datetime
import gspread
from google.oauth2.service_account import Credentials
import configparser
import sys
import tkinter as tk
from tkcalendar import Calendar
from openai import OpenAI
import re
# --- Load config from file ---
config = configparser.ConfigParser()
config.read("DONT_DELETE_ENV_FILES/config/00_STEP1_Nasean_Create_Essay_11_createMP3and_TimeStamp_short.txt")

ROOT_FOLDER = config.get("DEFAULT", "ROOT_FOLDER")
INPUT_FILE_NAME = config.get("DEFAULT", "INPUT_FILE")
OUTPUT_FOLDER_NAME = config.get("DEFAULT", "OUTPUT_FOLDER")
AUDIO_FILE_NAME = config.get("DEFAULT", "AUDIO_FILE")
selected_voice = config.get("DEFAULT", "VOICE")
tts_model = config.get("DEFAULT", "TTS_MODEL")
timestamp_file_name = config.get("DEFAULT", "TIMESTAMP_FILE")
whisper_model = config.get("DEFAULT", "WHISPER_MODEL")
max_chars = config.getint("DEFAULT", "MAX_TTS_CHARS")
retries = config.getint("DEFAULT", "RETRY_ATTEMPTS")
min_segment_duration = config.getint("DEFAULT", "MIN_SEGMENT_DURATION_SECONDS", fallback=5)

SERVICE_ACCOUNT_FILE = os.getenv("SMARTIKLE_WORKBOOK_GOOGLE_CREDENTIALS")
SHEET_ID = os.getenv("SMARTIKLE_GOOGLE_SHEET_ID")
api_key = os.getenv("OPENAI_API_KEY")

WORKSHEET_NAME = "G6VIR-short"

AUDIO_FILE_NAME = "narration_short.mp3"

if not api_key:
    print("❌ OPENAI_API_KEY not set.")
    sys.exit(1)

scope = [
    "https://www.googleapis.com/auth/spreadsheets.readonly",
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive"
]
creds = Credentials.from_service_account_file(SERVICE_ACCOUNT_FILE, scopes=scope)
client_gsheet = gspread.authorize(creds)
sheet = client_gsheet.open_by_key(SHEET_ID).worksheet(WORKSHEET_NAME)
rows = sheet.get_all_values()
headers = rows[0]
data_rows = rows[1:]

SYSTEM_INSTRUCTION_PATH = "DONT_DELETE_ENV_FILES/system_instructions.txt"
with open(SYSTEM_INSTRUCTION_PATH, "r", encoding="utf-8") as f:
    system_instruction = f.read().strip()

client = OpenAI(api_key=api_key)
session = requests.Session()

# --- Calendar date selector to find the starting row ---
selected_date = None

def on_submit():
    global selected_date
    selected_date = cal.get_date()
    root.destroy()

root = tk.Tk()
root.title("Select a Start Date")
today = datetime.today()
cal = Calendar(root, selectmode='day', year=today.year, month=today.month, day=today.day, date_pattern="yyyy-mm-dd")
cal.pack(pady=20)
tk.Button(root, text="Start", command=on_submit).pack(pady=10)
root.mainloop()

if not selected_date:
    print("❌ No date selected.")
    sys.exit(1)

print(f"📅 Selected date: {selected_date}")

# ✅ Save selected date to file for --use-date-file scripts
with open("selected_date.txt", "w") as f:
    f.write(selected_date)


# Convert calendar date to datetime for comparison
current_date = datetime.strptime(selected_date, "%Y-%m-%d")

# Find the first row that matches the selected calendar date in Column J
start_index = None
for idx, row in enumerate(data_rows, start=2):
    row_date = row[9].strip() if len(row) > 9 else ""
    if row_date == current_date.strftime("%Y-%m-%d"):
        start_index = idx
        break

if start_index is None:
    print(f"❌ No row found matching date {current_date.strftime('%Y-%m-%d')} in Column J")
    sys.exit(1)

# ✅ Process all rows from the matched row forward, using only status logic
for idx, row in enumerate(data_rows[start_index - 2:], start=start_index):
    row_date = row[9].strip() if len(row) > 9 else ""
    status = row[3].strip().lower() if len(row) > 3 else ""
    label = row[2].strip() if len(row) > 2 else ""

    print(f"🔍 Checking row {idx} — Date: {row_date}, Status: {status}, Label: {label}")

    if status == "stop creating":
        print(f"🛑 STOP CREATING flag found in row {idx}. Launching Step 12 script...")
        try:
            import subprocess
            subprocess.run(["python", "12_STEP2_Nasean_Generate_Image_from_prompts_short_V7.py", "--use-date-file"], check=True)
        except Exception as e:
            print(f"❌ Failed to run Step 12 script: {e}")
        break  # ✅ safely exit the loop after launching Step 12


    if status != "not started":
        print(f"⏭️ Skipping row {idx} — Status is '{status}'")
        continue

    print(f"🛠️ processing row {idx}...")
    sheet.update_cell(idx, 4, "in progress")
    try:
        parsed_date = datetime.strptime(row_date, "%Y-%m-%d")
        folder_name = f"{label}-{parsed_date.strftime('%m-%d-%Y')}"
    except ValueError:
        print(f"❌ Invalid date format in Column J for row {idx}: {row_date}")
        continue

    base_folder = Path(ROOT_FOLDER) / folder_name
    save_folder = base_folder / OUTPUT_FOLDER_NAME
    base_folder.mkdir(parents=True, exist_ok=True)
    save_folder.mkdir(parents=True, exist_ok=True)

    input_file = base_folder / INPUT_FILE_NAME
    audio_path = save_folder / AUDIO_FILE_NAME
    timestamp_path = save_folder / timestamp_file_name

    prompt_text = row[10].strip() if len(row) > 10 else ""
    if not input_file.exists():
        if prompt_text:
            print(f"📝 essay.docx not found for {folder_name}, generating from prompt...")
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": system_instruction},
                    {"role": "user", "content": prompt_text}
                ],
                temperature=0.7
            )
            reply = response.choices[0].message.content.strip()
            essay_only = reply.split("Multiple Choice Questions:")[0].strip()
            essay_only = essay_only.replace('\n', ' ')

            with open(save_folder / "response.txt", "w", encoding="utf-8") as f:
                f.write(reply)

            doc = docx.Document()
            doc.add_paragraph(essay_only)
            doc.save(input_file)
            doc.save(save_folder / INPUT_FILE_NAME)
            print(f"✅ essay.docx created for {folder_name}")
            sheet.update_cell(idx, 4, "document created")
        else:
            print(f"❌ No prompt in Column K and essay.docx does not exist. Skipping row.")
            continue
    else:
        print(f"📄 essay.docx already exists. Proceeding with narration.")

    story_lines = [p.text.strip() for p in docx.Document(input_file).paragraphs if p.text.strip()]
    full_story = "\n".join(story_lines)

    if not full_story.strip():
        print("❌ Story content is empty after cleaning.")
        continue

    print("🧠 Requesting YouTube title and description from OpenAI...")
    yt_prompt = f"""Based on the following educational essay, generate:
    1. A concise YouTube video title (max 100 characters)
    2. A short video description (2-3 lines) that summarizes the video content

    Essay:
    {full_story}
    """
    yt_response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a content creator skilled at writing compelling YouTube titles and descriptions."},
            {"role": "user", "content": yt_prompt}
        ],
        temperature=0.5
    )
    yt_result = yt_response.choices[0].message.content.strip()


    # Parse title and description cleanly
    title_match = re.search(r"1\.\s*YouTube Video Title:\s*\"?(.+?)\"?\s*$", yt_result, re.MULTILINE)
    desc_match = re.search(r"2\.\s*Video Description:\s*(.+)", yt_result, re.DOTALL)

    title_clean = title_match.group(1).strip() if title_match else "Untitled Video"
    desc_clean = desc_match.group(1).strip() if desc_match else "No description provided."

    # Write title and description to separate files
    (save_folder / "youtubetitle.txt").write_text(title_clean, encoding="utf-8")
    (save_folder / "youtubedescription.txt").write_text(desc_clean, encoding="utf-8")

    print("✅ YouTube title and description saved to separate files.")


    # 🌍 Extract essay metadata for image generation
    print("📡 Requesting historical metadata from OpenAI...")
    meta_prompt = f"""
    Based on the following essay, extract:
    1. The **time in history** when the event(s) took place (e.g., 1940s, Ancient Egypt, 19th century, etc.)
    2. The **geographical location** (e.g., country, city, region, or civilization)
    3. Any **relevant context or visual elements** useful for generating illustrations from this story (e.g., clothing styles, architecture, cultural elements, time of day, scenery, etc.)

    Please respond clearly in the following format:

    Time in history:
    Location:
    Visual context:
        
    Essay:
    {full_story}
    """

    meta_response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a historical researcher and visual storytelling assistant."},
            {"role": "user", "content": meta_prompt}
        ],
        temperature=0.5
    )

    meta_result = meta_response.choices[0].message.content.strip()
    meta_file = save_folder / "essay_metadata.txt"
    meta_file.write_text(meta_result, encoding="utf-8")
    print(f"🧾 Essay metadata saved to: {meta_file}")


    if not audio_path.exists():
        print("🎙️ Generating narration audio...")
        sheet.update_cell(idx, 4, "processing audio")        
        story_input = full_story[:max_chars]
        for attempt in range(retries):
            try:
                tts_response = session.post(
                    "https://api.openai.com/v1/audio/speech",
                    headers={
                        "Authorization": f"Bearer {api_key}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "model": tts_model,
                        "input": story_input,
                        "voice": selected_voice
                    },
                    timeout=60
                )
                if tts_response.status_code == 200:
                    with open(audio_path, "wb") as f:
                        f.write(tts_response.content)
                    print(f"✅ Narration saved to: {audio_path}")
                    break
                else:
                    print(f"⚠️ Attempt {attempt+1} failed: {tts_response.status_code}")
                    print(tts_response.text)
            except Exception as e:
                print(f"❌ TTS attempt {attempt+1} failed: {e}")
            time.sleep(5)
        else:
            print("❌ All TTS attempts failed. Skipping row.")
            continue

    print("🕒 Generating timestamped transcript...")
    model = whisper.load_model(whisper_model)
    result = model.transcribe(str(audio_path))
    with open(timestamp_path, "w", encoding="utf-8") as f:
        for seg in result['segments']:
            f.write(f"[{seg['start']:.2f}s - {seg['end']:.2f}s] {seg['text'].strip()}\n")

    if timestamp_path.exists() and timestamp_path.stat().st_size > 0:
        sheet.update_cell(idx, 4, "ready for images")
        print(f"✅ Timestamps saved for {folder_name}")
        
    else:
        sheet.update_cell(idx, 4, "error")
        print(f"❌ Failed to create timestamps for {folder_name}")
        # ✅ All rows processed — launch Step 12

import subprocess
print("🚀 Launching Step 12 image generation...")
subprocess.run(["python", "12_STEP2_Nasean_Generate_Image_from_prompts_short_V7.py", "--use-date-file"])
